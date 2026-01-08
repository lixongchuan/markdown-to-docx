import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt
from config import FONTS_MAP, ALIGNMENT_MAP, SPECIAL_CENTERED_KEYWORDS

class DocxGenerator:
    def __init__(self, markdown_text, config):
        self.md_text = markdown_text
        self.config = config
        self.doc = Document()
        self.md_parser = MarkdownIt("gfm-like", {'html': True})

    def generate(self):
        """执行生成流程"""
        self._setup_page_layout()
        self._define_styles()
        self._parse_and_build()
        return self.doc

    def _setup_page_layout(self):
        section = self.doc.sections[0]
        c = self.config
        section.top_margin = Inches(float(c.get('page_margin_top', 1.0)))
        section.bottom_margin = Inches(float(c.get('page_margin_bottom', 1.0)))
        section.left_margin = Inches(float(c.get('page_margin_left', 1.25)))
        section.right_margin = Inches(float(c.get('page_margin_right', 1.25)))
        
        if c.get('orientation') == 'LANDSCAPE':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT

    def _define_styles(self):
        """定义文档使用的所有样式"""
        c = self.config
        body_font = FONTS_MAP.get(c.get('body_font', 'song'), FONTS_MAP['song'])
        heading_font = FONTS_MAP.get(c.get('heading_font', 'hei'), FONTS_MAP['hei'])
        code_font = FONTS_MAP.get(c.get('code_font', 'courier'), FONTS_MAP['courier'])
        
        fs_pt = float(c.get('font_size', 12))
        line_spacing = float(c.get('line_spacing', 1.25))
        
        # 1. Normal (正文)
        style = self.doc.styles['Normal']
        self._set_font(style, body_font, fs_pt)
        pf = style.paragraph_format
        pf.alignment = ALIGNMENT_MAP.get(c.get('alignment', 'JUSTIFY'), WD_ALIGN_PARAGRAPH.JUSTIFY)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
        
        if c.get('indent_style') == 'first_line_2':
            pf.first_line_indent = Inches(0.3)
        
        pf.space_before = Pt(float(c.get('space_before', 0)))
        pf.space_after = Pt(float(c.get('space_after', 0)))

        # 2. Headings (标题)
        sizes = [16, 15, 14]
        for i, level in enumerate([1, 2, 3]):
            name = f'Final Heading {level}'
            style = self._get_or_create_style(name)
            self._set_font(style, heading_font, sizes[i], bold=True)
            style.paragraph_format.first_line_indent = Inches(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            style.paragraph_format.space_before = Pt(24 if level == 1 else 18)
            style.paragraph_format.space_after = Pt(18 if level == 1 else 12)

        # 3. Code Blocks
        style = self._get_or_create_style('CodeBlockText')
        self._set_font(style, code_font, max(9, fs_pt - 2))
        style.paragraph_format.first_line_indent = Inches(0)
        
        style = self._get_or_create_style('CodeBlockHeader')
        style.font.name = 'Arial'
        style.font.size = Pt(9)
        style.font.bold = True
        style.font.color.rgb = RGBColor(60, 60, 60)

        # 4. List Visuals
        style = self._get_or_create_style('ListParagraphVisual')
        style.base_style = self.doc.styles['Normal']
        style.paragraph_format.first_line_indent = Inches(0)

    def _parse_and_build(self):
        """解析 Markdown Token 并构建文档"""
        tokens = self.md_parser.parse(self.md_text)
        list_stack = []
        i = 0
        
        while i < len(tokens):
            token = tokens[i]
            
            # 列表状态管理
            if token.type in ('bullet_list_open', 'ordered_list_open'):
                start = token.attrs.get('start', 1) if token.attrs else 1
                list_stack.append({'type': 'ordered' if 'ordered' in token.type else 'bullet', 'index': int(start) - 1})
                i += 1; continue
            elif token.type in ('bullet_list_close', 'ordered_list_close'):
                if list_stack: list_stack.pop()
                i += 1; continue
            elif token.type == 'list_item_open':
                if list_stack: list_stack[-1]['index'] += 1
                i += 1; continue
            elif token.type == 'list_item_close':
                i += 1; continue

            # 内容处理
            if token.type == 'heading_open':
                level = int(token.tag[1:])
                self._add_heading(tokens[i+1], level)
                i += 3
            
            elif token.type == 'paragraph_open':
                self._add_paragraph(tokens[i+1], list_stack)
                i += 3
            
            elif token.type == 'table_open':
                i = self._process_table(tokens, i)
            
            elif token.type in ('fence', 'code_block'):
                self._add_code_block(token)
                i += 1
            else:
                i += 1

    # --- Helpers ---
    def _set_font(self, style, font_config, size_pt, bold=False):
        style.font.name = font_config['west']
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['east'])

    def _get_or_create_style(self, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
        if name not in self.doc.styles:
            return self.doc.styles.add_style(name, style_type)
        return self.doc.styles[name]

    def _process_inline(self, paragraph, inline_token):
        if not inline_token.children:
            if inline_token.content: paragraph.add_run(inline_token.content)
            return
        
        is_bold = False
        is_italic = False
        
        for child in inline_token.children:
            if child.type == 'text':
                run = paragraph.add_run(child.content)
                run.bold = is_bold
                run.italic = is_italic
            elif child.type == 'strong_open': is_bold = True
            elif child.type == 'strong_close': is_bold = False
            elif child.type == 'em_open': is_italic = True
            elif child.type == 'em_close': is_italic = False
            elif child.type == 'code_inline':
                run = paragraph.add_run(child.content)
                run.font.name = 'Courier New'

    def _add_heading(self, inline_token, level):
        style = f'Final Heading {level}' if 1 <= level <= 3 else 'Normal'
        p = self.doc.add_paragraph(style=style)
        content = inline_token.content.strip()
        if content in SPECIAL_CENTERED_KEYWORDS:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._process_inline(p, inline_token)

    def _add_paragraph(self, inline_token, list_stack):
        if list_stack:
            p = self.doc.add_paragraph(style='ListParagraphVisual')
            self._handle_list_indent(p, list_stack)
        else:
            p = self.doc.add_paragraph()
            if inline_token.content.strip() in SPECIAL_CENTERED_KEYWORDS:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._process_inline(p, inline_token)

    def _handle_list_indent(self, paragraph, list_stack):
        level = len(list_stack) - 1
        current = list_stack[-1]
        
        base = 0.15 if self.config.get('indent_style') == 'first_line_2' else 0.0
        left_indent = base + 0.22 + (level * 0.22)
        
        paragraph.paragraph_format.left_indent = Inches(left_indent)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(left_indent), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        
        prefix = ["●", "○", "■"][level % 3] if current['type'] == 'bullet' else f"{current['index']}."
        paragraph.add_run(f"{prefix}\t").font.name = 'Times New Roman'

    def _add_code_block(self, token):
        table = self.doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        
        # Header
        cell = table.cell(0, 0)
        self._set_shading(cell, 'D9D9D9')
        p = cell.paragraphs[0]
        p.style = 'CodeBlockHeader'
        p.add_run(token.info.strip() if token.info else "Code")
        
        # Content
        cell = table.cell(1, 0)
        self._set_shading(cell, 'F2F2F2')
        p = cell.paragraphs[0]
        p.style = 'CodeBlockText'
        p.text = token.content.rstrip()

    def _process_table(self, tokens, start_index):
        # 简化版表格处理逻辑
        # 实际代码中需要遍历 tokens 直到 table_close
        # 这里为了演示简洁省略部分遍历逻辑，实际可参考原代码迁移
        # 寻找 table_close
        j = start_index + 1
        rows_data = []
        current_row = []
        
        while j < len(tokens) and tokens[j].type != 'table_close':
            tok = tokens[j]
            if tok.type == 'tr_open': current_row = []
            elif tok.type in ('th_open', 'td_open'):
                if j+1 < len(tokens) and tokens[j+1].type == 'inline':
                    current_row.append(tokens[j+1])
            elif tok.type == 'tr_close':
                rows_data.append(current_row)
            j += 1
        
        if rows_data:
            table = self.doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r, row in enumerate(rows_data):
                for c, token in enumerate(row):
                    if c < len(table.columns):
                        p = table.cell(r, c).paragraphs[0]
                        self._process_inline(p, token)
        
        return j + 1

    def _set_shading(self, cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tc_pr.append(shd)