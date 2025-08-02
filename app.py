# app.py (V37.11 - 代码块空行修正版)
import io, datetime, re, requests, uuid
from flask import Flask, request, send_file, render_template, jsonify
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt

app = Flask(__name__)
md = MarkdownIt("gfm-like")

# --- 样式定义 (新增HangingParagraph样式) ---
def define_final_styles(doc):
    # --- 正文、标题、代码、表格样式 ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p_format = style.paragraph_format; p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_format.first_line_indent = Inches(0.25 * 1.2) # 正文保留首行缩进
    p_format.space_before = Pt(0); p_format.space_after = Pt(0)
    
    style = doc.styles.add_style('Final Heading 1', 1)
    style.font.name = 'Times New Roman'; style.font.size = Pt(16); style.font.bold = True
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    p_format = style.paragraph_format; p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_format.space_before = Pt(24); p_format.space_after = Pt(12)

    style = doc.styles.add_style('Final Heading 2', 1)
    style.font.name = 'Times New Roman'; style.font.size = Pt(14); style.font.bold = True
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    p_format = style.paragraph_format; p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format.space_before = Pt(12); p_format.space_after = Pt(12)
    
    style = doc.styles.add_style('Final Heading 3', 1)
    style.font.name = 'Times New Roman'; style.font.size = Pt(12); style.font.bold = True
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    p_format = style.paragraph_format; p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format.space_before = Pt(6); p_format.space_after = Pt(6)

    try:
        style = doc.styles.add_style('CodeStyle', 1)
        style.font.name = 'Courier New'; style.font.size = Pt(10)
        p_format = style.paragraph_format; p_format.space_before = Pt(6); p_format.space_after = Pt(6)
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F1F1F1')
        style.paragraph_format.element.get_or_add_pPr().append(shd)
    except ValueError: pass
    
    try:
        style = doc.styles.add_style('TableCellClean', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        p_format = style.paragraph_format
        p_format.first_line_indent = Inches(0)
        p_format.left_indent = Inches(0)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except ValueError:
        pass
    
    # V37.10 新增: 为以'-'开头的特殊段落创建无任何缩进的专属样式
    try:
        style = doc.styles.add_style('HangingParagraph', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.base_style = doc.styles['Normal'] # 基于正文样式，但覆盖缩进
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        p_format = style.paragraph_format
        p_format.first_line_indent = Inches(0) # 关键: 无首行缩进
        p_format.left_indent = Inches(0)      # 关键: 无左侧缩进
    except ValueError:
        pass

# --- 辅助函数 ---
def add_pageref_field(paragraph, bookmark_name):
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar'); fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(nsdecls('w'), 'preserve'); instrText.text = f' PAGEREF {bookmark_name} \\h '
    fldChar_separate = OxmlElement('w:fldChar'); fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run_text = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '1'; run_text.append(t)
    fldChar_end = OxmlElement('w:fldChar'); fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin); run._r.append(instrText); run._r.append(fldChar_separate); run._r.append(run_text); run._r.append(fldChar_end)

def create_bookmark(paragraph, name):
    bookmark_start = OxmlElement('w:bookmarkStart'); bookmark_start.set(qn('w:id'), '0'); bookmark_start.set(qn('w:name'), name)
    paragraph._p.insert(0, bookmark_start)
    bookmark_end = OxmlElement('w:bookmarkEnd'); bookmark_end.set(qn('w:id'), '0')
    paragraph._p.append(bookmark_end)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part; r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr'); c_style = OxmlElement('w:rStyle'); c_style.set(qn('w:val'), 'Hyperlink')
    rPr.append(c_style); new_run.append(rPr); text_element = OxmlElement('w:t'); text_element.text = text; new_run.append(text_element)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)

def _process_inline(paragraph, inline_token):
    """
    这个函数负责处理一行文字中的所有内联格式（加粗、斜体、链接等）。
    它遍历子令牌，而不是直接使用原始文本，从而“消耗掉”格式化符号。
    """
    style_stack = {'bold': False, 'italic': False}
    if not inline_token.children:
        # 对于没有内联格式的简单文本
        paragraph.add_run(inline_token.content)
        return

    for child in inline_token.children:
        if child.type == 'text':
            run = paragraph.add_run(child.content)
            run.bold = style_stack['bold']
            run.italic = style_stack['italic']
        elif child.type == 'strong_open': style_stack['bold'] = True
        elif child.type == 'strong_close': style_stack['bold'] = False
        elif child.type == 'em_open': style_stack['italic'] = True
        elif child.type == 'em_close': style_stack['italic'] = False
        elif child.type == 'link_open':
            url = child.attrs.get('href', '')
            # 链接文本通常在link_open的第一个子节点中
            link_text = ""
            if child.children and child.children[0].type == 'text':
                link_text = child.children[0].content
            else:
                # 如果没有文本子节点，使用URL作为文本
                link_text = url
            
            if not url.startswith('#'): # 外部链接
                add_hyperlink(paragraph, link_text, url)
            else: # 内部书签链接
                paragraph.add_run(link_text) # 暂时只添加文本，不处理内部链接跳转
        # 其他如 code_inline, image 等可以按需添加

def _trim_inline_token_whitespace(inline_token):
    if not inline_token.children:
        if inline_token.content:
             inline_token.content = inline_token.content.strip()
        return

    for child in inline_token.children:
        if child.type == 'text':
            child.content = child.content.lstrip()
            break
            
    for child in reversed(inline_token.children):
        if child.type == 'text':
            child.content = child.content.rstrip()
            break

# --- 核心转换引擎 (V37.11 修正代码块处理) ---
def markdown_to_docx(md_text: str):
    document = Document()
    define_final_styles(document)
    tokens = md.parse(md_text)
    
    SPECIAL_CENTERED_KEYWORDS = {"附录", "引言", "目录", "摘要", "参考文献"}
    bookmark_map = {}
    is_toc_present = "目录" in md_text
    toc_heading_text = "目录"
    
    # 预先收集所有标题的书签信息
    for i, token in enumerate(tokens):
        if token.type == 'heading_open':
            # 确保tokens[i+1]是inline类型且有content属性
            if i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                content = tokens[i+1].content.strip()
                if content: # 只有非空标题才创建书签
                    bookmark_map[content] = f"bkmk_{uuid.uuid4().hex[:8]}"

    i = 0
    while i < len(tokens):
        token = tokens[i]
        
        # 检查是否为需要居中的特殊标题或段落
        is_special_centered_heading = False
        is_special_centered_paragraph = False
        if token.type == 'heading_open' and i + 1 < len(tokens) and tokens[i+1].type == 'inline' and tokens[i+1].content.strip() in SPECIAL_CENTERED_KEYWORDS:
             is_special_centered_heading = True
        # 注意：这里判断is_special_centered_paragraph时，要确保tokens[i+1]是inline类型
        if token.type == 'paragraph_open' and i + 1 < len(tokens) and tokens[i+1].type == 'inline' and tokens[i+1].content.strip() in SPECIAL_CENTERED_KEYWORDS:
             is_special_centered_paragraph = True

        # 目录处理逻辑 (V37.7 保持不变)
        if is_toc_present and token.type == 'heading_open' and i + 1 < len(tokens) and tokens[i+1].type == 'inline' and tokens[i+1].content.strip() == toc_heading_text:
            p_toc_heading = document.add_paragraph(toc_heading_text, style='Final Heading 1')
            p_toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # 确保目录标题居中
            tab_stops = p_toc_heading.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
            
            # 查找目录列表的起始位置
            list_start_index = i + 3 # 跳过 heading_open, inline, heading_close
            
            # 确保下一个令牌是列表开始
            if list_start_index < len(tokens) and tokens[list_start_index].type == 'bullet_list_open':
                j = list_start_index + 1
                while j < len(tokens) and tokens[j].type != 'bullet_list_close':
                    if tokens[j].type == 'list_item_open':
                        # 目录项的结构通常是 list_item_open -> paragraph_open -> inline
                        if j + 2 < len(tokens) and tokens[j+1].type == 'paragraph_open' and tokens[j+2].type == 'inline':
                            # 提取目录项文本
                            # 这里假设目录项文本是inline token的直接内容，不包含复杂的内联格式
                            # 如果目录项本身包含加粗等，需要修改此处以调用_process_inline
                            text = tokens[j+2].content.strip() 
                            
                            p_toc = document.add_paragraph()
                            p_toc.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
                            p_toc.add_run(text)
                            p_toc.add_run('\t')
                            
                            bookmark_name = bookmark_map.get(text)
                            if bookmark_name: 
                                add_pageref_field(p_toc, bookmark_name)
                        
                        # 移动 j 到当前 list_item_close 之后
                        k = j + 1
                        while k < len(tokens) and tokens[k].type != 'list_item_close':
                            k += 1
                        j = k # j 现在指向 list_item_close
                    j += 1 # 移动到下一个令牌
                i = j + 1 # 跳过整个目录列表
                continue
            else: # 如果目录标题后没有紧跟列表，则只处理标题
                i += 3 
                continue

        # V37.8 修正的标题处理逻辑 (保留)
        elif token.type == 'heading_open':
            level = int(token.tag[1:])
            style_name = f'Final Heading {level}' if 1 <= level <= 3 else 'Normal'
            
            p = document.add_paragraph(style=style_name)
            
            # 确保 inline_token 存在且类型正确
            if i + 1 < len(tokens) and tokens[i+1].type == 'inline':
                inline_token = tokens[i+1]
                _process_inline(p, inline_token)
                
                content_text = inline_token.content.strip()
                if content_text in SPECIAL_CENTERED_KEYWORDS:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                bookmark_name = bookmark_map.get(content_text)
                if bookmark_name: create_bookmark(p, bookmark_name)
            
            i += 3 # 跳过 heading_open, inline, heading_close
            continue
            
        # V37.9 核心修改点: 删除了整个 `elif token.type == 'bullet_list_open':` 代码块
        # 这使得列表项会被下面的 `paragraph_open` 逻辑当作普通段落处理，从而消除黑点

        # 普通段落处理逻辑 (现在也会处理之前的列表项)
        elif token.type == 'paragraph_open':
            inline_token = tokens[i+1]
            content_text = inline_token.content.strip()

            # 判断是否为特殊居中段落
            if content_text in SPECIAL_CENTERED_KEYWORDS:
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _process_inline(p, inline_token)
            
            # V37.10 新增: 判断是否为悬挂段落
            elif content_text.startswith('- '):
                p = document.add_paragraph(style='HangingParagraph')
                # 我们需要移除开头的 '- '，因为它只是一个格式标记
                # 智能地处理，只移除第一个文本节点的 '- '
                if inline_token.children and inline_token.children[0].type == 'text':
                    inline_token.children[0].content = inline_token.children[0].content.lstrip('- ')
                _process_inline(p, inline_token)
            
            # 其他所有普通段落
            else:
                p = document.add_paragraph() # 默认使用 'Normal' 样式
                _process_inline(p, inline_token)
                
            i += 3
            continue

        # 表格处理逻辑 (使用 V37.7 的可靠版本，保留)
        elif token.type == 'table_open':
            # V37.7: 全新、极简且绝对可靠的表格数据提取逻辑
            table_data = {'header': [], 'body': []}
            current_row = []
            in_header = False

            j = i + 1
            while j < len(tokens) and tokens[j].type != 'table_close':
                tok = tokens[j]
                if tok.type == 'thead_open': in_header = True
                elif tok.type == 'thead_close': in_header = False
                elif tok.type == 'tr_open': current_row = []
                elif tok.type in ('th_open', 'td_open'):
                    # 确保 inline_content 存在且类型正确
                    if j + 1 < len(tokens) and tokens[j+1].type == 'inline':
                        inline_content = tokens[j+1]
                        _trim_inline_token_whitespace(inline_content)
                        current_row.append(inline_content)
                elif tok.type == 'tr_close':
                    if in_header:
                        table_data['header'] = current_row
                    else:
                        table_data['body'].append(current_row)
                j += 1
            i = j # 主循环跳过整个表格
            
            # --- V37.7: 最核心的修改：创建表格并强制居中所有单元格 ---
            header_content = table_data['header']
            body_rows = table_data['body']

            if header_content:
                num_cols = len(header_content)
                num_rows = 1 + len(body_rows)
                table = document.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # 统一的填充函数
                def fill_and_force_center(cell, inline_token):
                    # 步骤一：获取段落并应用无缩进的干净样式
                    p = cell.paragraphs[0]
                    p.style = 'TableCellClean'
                    
                    # 步骤二：填充内容
                    _process_inline(p, inline_token)
                    
                    # 步骤三：强制将段落居中 (遵照您的指示)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 步骤四：设置垂直居中
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # 填充表头
                for col_idx, inline_token in enumerate(header_content):
                    fill_and_force_center(table.rows[0].cells[col_idx], inline_token)
                
                # 填充表体
                for row_idx, row_data in enumerate(body_rows):
                    for col_idx, inline_token in enumerate(row_data):
                        fill_and_force_center(table.rows[row_idx + 1].cells[col_idx], inline_token)

            continue
            
        # V37.11 核心修正点: 代码块处理
        if token.type == 'fence':
            # 使用 rstrip() 清理掉代码内容末尾的所有空白字符（包括换行符）
            cleaned_content = token.content.rstrip()
            document.add_paragraph(cleaned_content, style='CodeStyle')
            i += 1
            continue

        # 其他所有令牌
        else:
            i += 1
            
    return document

# --- API 路由 (无变化) ---
@app.route('/')
def index(): return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    try:
        markdown_text = request.form['markdown']
        if not markdown_text.strip(): return jsonify({'error': '内容不能为空.'}), 400
        document_obj = markdown_to_docx(markdown_text)
        file_stream = io.BytesIO()
        document_obj.save(file_stream)
        file_stream.seek(0)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        return send_file(
            file_stream, as_attachment=True, download_name=f'Markdown_Doc_{timestamp}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        import traceback; traceback.print_exc() 
        return jsonify({'error': f'服务器内部错误: {e}'}), 500

if __name__ == '__main__': app.run(host='0.0.0.0', port=5000, debug=True)
